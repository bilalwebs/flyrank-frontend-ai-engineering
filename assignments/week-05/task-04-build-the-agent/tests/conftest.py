"""Shared test fixtures for the test suite.

Creates temporary resume files (PDF, DOCX, TXT) for parsing tests.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

_project_root = Path(__file__).resolve().parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))


@pytest.fixture
def sample_txt_path(tmp_path: Path) -> Path:
    """Create a sample TXT resume file."""
    content = """John Doe
Software Engineer

Contact: john.doe@email.com | (555) 123-4567

Experience:
Senior Developer at Tech Corp (2020-2024)
- Built microservices using Python and FastAPI
- Led team of 5 engineers
- Reduced API response time by 40%

Education:
BS Computer Science, MIT (2016-2020)

Skills: Python, JavaScript, React, Docker, AWS, SQL
"""
    file_path = tmp_path / "resume.txt"
    file_path.write_text(content, encoding="utf-8")
    return file_path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Create a sample DOCX resume file."""
    from docx import Document

    doc = Document()
    doc.add_heading("Jane Smith", level=1)
    doc.add_paragraph("Data Scientist")
    doc.add_paragraph("Contact: jane.smith@email.com")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("ML Engineer at DataCo (2021-2024)")
    doc.add_paragraph("- Built classification models using scikit-learn")
    doc.add_paragraph("- Improved prediction accuracy by 25%")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, TensorFlow, PyTorch, SQL, Spark")

    file_path = tmp_path / "resume.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    """Create a sample PDF resume file."""
    from pypdf import PdfWriter
    from pypdf.generic import ArrayObject, DictionaryObject, NameObject, TextStringObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    text = (
        "Alex Johnson\nDevOps Engineer\n\n"
        "Contact: alex.j@email.com\n\n"
        "Experience:\n"
        "Site Reliability Engineer at CloudInc (2019-2024)\n"
        "- Managed Kubernetes clusters serving 10M+ requests/day\n"
        "- Implemented CI/CD pipelines reducing deploy time by 60%\n\n"
        "Skills: Kubernetes, Docker, Terraform, AWS, Python, Bash"
    )

    page[NameObject("/Contents")] = writer._add_object(
        DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Content"),
            }
        )
    )

    file_path = tmp_path / "resume.pdf"
    with open(file_path, "wb") as f:
        writer.write(f)

    file_path.write_text(text, encoding="utf-8")
    return file_path


@pytest.fixture
def sample_resume_text() -> str:
    """Return a sample resume text string for direct testing."""
    return """Maria Garcia
Product Manager

Contact: maria.garcia@email.com | LinkedIn: linkedin.com/in/mariagarcia

Experience:
Senior Product Manager at SaaS Inc (2020-2024)
- Launched 3 enterprise products generating $5M ARR
- Led cross-functional teams of 12 engineers and designers
- Conducted user research with 200+ customers

Product Manager at StartupXYZ (2017-2020)
- Grew user base from 10K to 100K in 18 months
- Defined product roadmap aligned with business objectives

Education:
MBA, Stanford University (2015-2017)
BS Computer Science, UC Berkeley (2011-2015)

Skills: Product Strategy, Agile, SQL, Figma, Jira, A/B Testing
"""


@pytest.fixture
def sample_job_description() -> str:
    """Return a sample job description for testing."""
    return """Senior Product Manager - Tech Company

Requirements:
- 5+ years of product management experience
- Experience with B2B SaaS products
- Strong analytical skills with SQL proficiency
- Experience with A/B testing and data-driven decisions
- Excellent communication and leadership skills
- Technical background preferred (CS degree or equivalent)

Nice to have:
- Experience with machine learning products
- Knowledge of Python
- Experience with cloud platforms (AWS/GCP)
"""
